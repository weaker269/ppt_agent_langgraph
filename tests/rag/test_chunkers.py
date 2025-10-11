"""验证文档加载与分块逻辑"""

from __future__ import annotations

from pathlib import Path

import fitz  # type: ignore
from docx import Document as DocxDocument  # type: ignore
from docx.enum.style import WD_STYLE_TYPE  # type: ignore

from src.rag.chunkers import chunk_document
from src.rag.loaders import load_document
from src.rag.models import DocumentMetadata, DocumentSection, LoadedDocument


def test_markdown_chunking_produces_metadata() -> None:
    path = Path("sample.md")
    loaded = load_document(path)
    chunks = chunk_document(loaded, chunk_size=220, sentence_overlap=1)

    assert chunks, "Markdown 文档应拆分出至少一个 chunk"
    first_chunk = chunks[0]
    assert first_chunk.section_title is not None
    assert first_chunk.length <= 260
    assert first_chunk.start_char >= 0


def test_plain_text_loader_and_chunking(tmp_path) -> None:
    content = (
        "这是一个包含多句的纯文本示例，用于验证 fallback 逻辑，"
        "确保没有标题时也能完成分段，同时保持内容完整。"
    )
    file_path = tmp_path / "note.txt"
    file_path.write_text(content, encoding="utf-8")

    loaded = load_document(file_path)
    chunks = chunk_document(loaded, chunk_size=80, sentence_overlap=0)

    assert chunks
    assert all(chunk.section_title is None for chunk in chunks)
    assert max(chunk.length for chunk in chunks) <= 120


def test_docx_loader_recognizes_headings(tmp_path) -> None:
    docx_path = tmp_path / "demo.docx"
    document = DocxDocument()
    document.add_heading("第一章 总体概述", level=1)
    document.add_paragraph("这一小节应该标注在章标题之下。")
    document.add_heading("1.1 目标", level=2)
    document.add_paragraph("目标细分内容。")
    document.save(docx_path)

    loaded = load_document(docx_path)
    sections = loaded.sections
    assert any(section.title == "第一章 总体概述" for section in sections)
    assert any(section.title == "1.1 目标" for section in sections)

    chunks = chunk_document(loaded, chunk_size=120, sentence_overlap=1)
    assert chunks
    assert chunks[0].section_title == "第一章 总体概述"


def test_docx_loader_recognizes_localized_heading(tmp_path) -> None:
    docx_path = tmp_path / "demo_cn_style.docx"
    document = DocxDocument()
    styles = document.styles
    try:
        cn_style = styles["标题 1"]  # type: ignore[index]
    except KeyError:
        cn_style = styles.add_style("标题 1", WD_STYLE_TYPE.PARAGRAPH)  # type: ignore[arg-type]
        cn_style.base_style = styles["Heading 1"]

    paragraph = document.add_paragraph("本地化标题")
    paragraph.style = cn_style
    document.add_paragraph("普通段落内容")
    document.save(docx_path)

    loaded = load_document(docx_path)
    assert any(section.title == "本地化标题" for section in loaded.sections)


def test_pdf_loader_preserves_page_numbers(tmp_path) -> None:
    pdf_path = tmp_path / "demo.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "第一页内容。第二句。")
    doc.save(pdf_path)
    doc.close()

    loaded = load_document(pdf_path)
    chunks = chunk_document(loaded, chunk_size=120, sentence_overlap=1)

    assert chunks
    assert chunks[0].page_number == 1
    assert chunks[0].section_title == "Page 1"


def test_chunk_metadata_is_isolated() -> None:
    metadata = DocumentMetadata(
        document_id="doc-1",
        source_path="doc.md",
        media_type="text/markdown",
    )
    section = DocumentSection(
        section_id="sec-1",
        title="第一节",
        level=1,
        text="这是一段内容。",
        metadata={"topic": "原始"},
    )
    document = LoadedDocument(metadata=metadata, sections=[section], full_text="这是一段内容。")

    chunks = chunk_document(document, chunk_size=120, sentence_overlap=0)
    assert chunks
    chunks[0].metadata["topic"] = "修改"

    assert document.sections[0].metadata["topic"] == "原始"
