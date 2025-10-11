"""文档加载器，负责解析多种格式。"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Callable, Iterable, List, Optional

import fitz  # type: ignore
from docx import Document as DocxDocument  # type: ignore

from .models import DocumentMetadata, DocumentSection, LoadedDocument


LoaderFunc = Callable[[Path], LoadedDocument]


def _sanitize_text(text: str) -> str:
    """统一去除 BOM、回车等无关字符。"""

    if text.startswith("\ufeff"):
        text = text.lstrip("\ufeff")
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _build_metadata(path: Path, media_type: str) -> DocumentMetadata:
    return DocumentMetadata(
        document_id=path.stem,
        source_path=str(path),
        media_type=media_type,
    )


def _markdown_sections(text: str) -> List[DocumentSection]:
    sections: List[DocumentSection] = []
    current_lines: List[str] = []
    current_title: Optional[str] = None
    current_level = 1
    current_start = 0
    cursor = 0
    section_index = 0

    heading_pattern = re.compile(r"^(#{1,6})\s+(.*)")
    lines = text.splitlines(keepends=True)

    for line in lines:
        line_length = len(line)
        match = heading_pattern.match(line.strip())
        if match:
            # 先结束上一节
            if current_lines:
                section_text = "".join(current_lines).strip()
                sections.append(
                    DocumentSection(
                        section_id=f"sec_{section_index:03d}",
                        title=current_title,
                        level=current_level,
                        text=section_text,
                        start_char=current_start,
                        end_char=cursor,
                    )
                )
                section_index += 1
                current_lines = []

            current_level = len(match.group(1))
            current_title = match.group(2).strip()
            current_start = cursor + line_length
        else:
            if not current_lines:
                current_start = cursor
            current_lines.append(line)
        cursor += line_length

    if current_lines:
        section_text = "".join(current_lines).strip()
        sections.append(
            DocumentSection(
                section_id=f"sec_{section_index:03d}",
                title=current_title,
                level=current_level,
                text=section_text,
                start_char=current_start,
                end_char=cursor,
            )
        )

    # 若文本在第一节前有内容，则补充一个默认段落
    if not sections:
        sections.append(
            DocumentSection(
                section_id="sec_000",
                title=None,
                level=1,
                text=text.strip(),
                start_char=0,
                end_char=len(text),
            )
        )

    return sections


def load_markdown(path: Path) -> LoadedDocument:
    raw_text = path.read_text(encoding="utf-8")
    text = _sanitize_text(raw_text)
    metadata = _build_metadata(path, "text/markdown")
    sections = _markdown_sections(text)
    return LoadedDocument(metadata=metadata, sections=sections, full_text=text)


def load_plain_text(path: Path) -> LoadedDocument:
    raw_text = path.read_text(encoding="utf-8")
    text = _sanitize_text(raw_text)
    metadata = _build_metadata(path, "text/plain")
    section = DocumentSection(
        section_id="sec_000",
        title=None,
        level=1,
        text=text,
        start_char=0,
        end_char=len(text),
    )
    return LoadedDocument(metadata=metadata, sections=[section], full_text=text)


def load_pdf(path: Path) -> LoadedDocument:
    doc = fitz.open(path)
    sections: List[DocumentSection] = []
    offset = 0
    buffer = io.StringIO()
    for index, page in enumerate(doc, start=1):
        page_text = page.get_text("text")
        clean_text = _sanitize_text(page_text)
        page_start = offset
        page_end = page_start + len(clean_text)
        buffer.write(clean_text)
        sections.append(
            DocumentSection(
                section_id=f"page_{index:03d}",
                title=f"Page {index}",
                level=1,
                text=clean_text,
                start_char=page_start,
                end_char=page_end,
                page_number=index,
            )
        )
        offset = page_end
    doc.close()
    metadata = _build_metadata(path, "application/pdf")
    return LoadedDocument(
        metadata=metadata,
        sections=sections,
        full_text=buffer.getvalue(),
    )


def load_docx(path: Path) -> LoadedDocument:
    document = DocxDocument(path)
    sections: List[DocumentSection] = []
    buffer = io.StringIO()
    offset = 0
    section_index = 0
    current_title: Optional[str] = None
    current_level = 1
    current_start = 0
    current_lines: List[str] = []

    def flush_section(final_offset: int) -> None:
        nonlocal section_index, current_lines, current_title, current_start
        if not current_lines:
            return
        text = "\n".join(current_lines).strip()
        sections.append(
            DocumentSection(
                section_id=f"sec_{section_index:03d}",
                title=current_title,
                level=current_level,
                text=text,
                start_char=current_start,
                end_char=final_offset,
            )
        )
        section_index += 1
        current_lines = []

    for paragraph in document.paragraphs:
        para_text = paragraph.text.strip()
        if not para_text:
            continue

        buffer.write(para_text + "\n")
        new_offset = offset + len(para_text) + 1

        style_name = paragraph.style.name if paragraph.style else "Normal"
        style_name_lower = style_name.lower()
        is_heading = style_name_lower.startswith("heading") or "标题" in style_name
        if is_heading:
            flush_section(offset)
            heading_match = re.search(r"(\d+)", style_name)
            level = int(heading_match.group(1)) if heading_match else 1
            current_title = para_text
            current_level = level
            current_start = offset
        else:
            if not current_lines:
                current_start = offset
            current_lines.append(para_text)

        offset = new_offset

    flush_section(offset)

    if not sections and buffer.tell():
        text = buffer.getvalue().strip()
        sections.append(
            DocumentSection(
                section_id="sec_000",
                title=None,
                level=1,
                text=text,
                start_char=0,
                end_char=len(text),
            )
        )

    metadata = _build_metadata(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return LoadedDocument(
        metadata=metadata,
        sections=sections,
        full_text=buffer.getvalue().strip(),
    )


def _choose_loader(path: Path) -> LoaderFunc:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        return load_markdown
    if suffix in {".txt"}:
        return load_plain_text
    if suffix in {".pdf"}:
        return load_pdf
    if suffix in {".docx"}:
        return load_docx
    raise ValueError(f"暂不支持的文档类型: {suffix}")


def load_document(path: Path) -> LoadedDocument:
    loader = _choose_loader(path)
    return loader(path)


def load_documents(paths: Iterable[Path]) -> List[LoadedDocument]:
    return [load_document(Path(p)) for p in paths]


__all__ = [
    "load_document",
    "load_documents",
    "load_markdown",
    "load_plain_text",
    "load_pdf",
    "load_docx",
]
